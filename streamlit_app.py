import streamlit as st
import requests
from datetime import datetime, timedelta
import pandas as pd
import numpy as np
import io
from docx import Document

# --- CONFIGURACIÓN ---
try:
    ATHLETE_ID = st.secrets["ATHLETE_ID"]
    API_KEY = st.secrets["API_KEY"]
except (FileNotFoundError, KeyError):
    ATHLETE_ID = "i10474"
    API_KEY = "27i9azt55smmhvg1ogc5gmn7x"

# --- FUNCIONES AUXILIARES ---
def format_duration(seconds):
    if not isinstance(seconds, (int, float)) or seconds < 0: return "0m"
    h, m = divmod(seconds // 60, 60)
    return f"{int(h)}h {int(m)}m" if h > 0 else f"{int(m)}m"

def dataframe_to_word(df):
    """Convierte un DataFrame de Pandas a un documento de Word en memoria."""
    if 'Fecha' in df.columns:
        df['Fecha'] = pd.to_datetime(df['Fecha']).dt.strftime('%d-%m-%Y')
        
    doc = Document()
    doc.add_heading('Resumen de Actividades', level=1)
    
    table = doc.add_table(rows=1, cols=df.shape[1])
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for j, col_name in enumerate(df.columns):
        hdr_cells[j].text = str(col_name)
        hdr_cells[j].paragraphs[0].runs[0].font.bold = True

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, col_name in enumerate(df.columns):
            row_cells[j].text = str(row[col_name])

    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    return mem_file

# --- FUNCIÓN DE READINESS ---
@st.cache_data(ttl=3600)
def get_readiness_analysis_v3(selected_date, api_key, athlete_id):
    start_date = selected_date - timedelta(days=60)
    end_date = selected_date
    params = {'oldest': start_date.strftime('%Y-%m-%d'), 'newest': end_date.strftime('%Y-%m-%d')}
    wellness_url = f"https://intervals.icu/api/v1/athlete/{athlete_id}/wellness"
    
    try:
        response = requests.get(wellness_url, auth=('API_KEY', api_key), params=params)
        if response.status_code != 200 or not response.json():
            return {"error": "No se encontraron suficientes datos de bienestar."}
    except requests.exceptions.RequestException as e:
        return {"error": f"Error de conexión: {e}"}

    df = pd.DataFrame(response.json())
    df['id'] = pd.to_datetime(df['id'])
    df.set_index('id', inplace=True)
    df = df.sort_index()
    
    today_str = selected_date.strftime('%Y-%m-%d')
    if pd.to_datetime(today_str) not in df.index:
        return {"error": f"No hay datos de bienestar para el día {selected_date.strftime('%d-%m-%Y')}"}

    today_data = df.loc[today_str]
    past_df = df.loc[df.index < pd.to_datetime(today_str)]
    hrv_baseline_28d = past_df['hrv'].tail(28).mean()
    hrv_std_28d = past_df['hrv'].tail(28).std()
    hrv_hoy, rhr_hoy, sleep_score_hoy = today_data.get('hrv'), today_data.get('restingHR'), today_data.get('sleepScore')

    score, breakdown = 0, []
    if pd.notna(hrv_hoy) and pd.notna(hrv_baseline_28d) and pd.notna(hrv_std_28d):
        hrv_points, hrv_normal_range_lower = 0, hrv_baseline_28d - (0.75 * hrv_std_28d)
        if hrv_hoy >= hrv_baseline_28d + (0.5 * hrv_std_28d): hrv_points = 45
        elif hrv_hoy >= hrv_normal_range_lower: hrv_points = 30
        elif hrv_hoy >= hrv_baseline_28d - hrv_std_28d: hrv_points = 15
        score += hrv_points
        breakdown.append(f"**VFC (HRV):** `{hrv_hoy:.1f}ms`. Rango normal: `{hrv_normal_range_lower:.1f}ms - {hrv_baseline_28d + (0.5 * hrv_std_28d):.1f}ms` → **{hrv_points} ptos**.")
    if pd.notna(rhr_hoy):
        rhr_points = 0
        if rhr_hoy <= 45: rhr_points = 35
        elif rhr_hoy <= 48: rhr_points = 25
        elif rhr_hoy <= 52: rhr_points = 10
        score += rhr_points
        breakdown.append(f"**FC Reposo:** `{rhr_hoy:.0f}bpm` → **{rhr_points} ptos**.")
    if pd.notna(sleep_score_hoy):
        sleep_points = 20 if sleep_score_hoy >= 80 else 10 if sleep_score_hoy >= 70 else 0
        score += sleep_points
        breakdown.append(f"**P. Sueño:** `{sleep_score_hoy:.0f}` → **{sleep_points} ptos**.")
    
    verdict_text = "🚫 **LUZ ROJA:** Recuperación prioritaria."
    if score >= 80: verdict_text = "✅ **LUZ VERDE:** Estado óptimo."
    elif score >= 60: verdict_text = "⚠️ **LUZ AMARILLA:** Procede con cautela."
    return {"verdict": verdict_text, "readiness_score": score, "score_breakdown": breakdown}

def display_gauge(score):
    score_color = "#d9534f" if score < 60 else "#f0ad4e" if score < 80 else "#5cb85c"
    st.markdown(f"""<div style="background-color: #f1f1f1; border-radius: 5px; padding: 2px;"><div style="background-color: {score_color}; width: {score}%; height: 24px; border-radius: 5px; text-align: center; color: white; font-weight: bold; line-height: 24px;">{score} / 100</div></div>""", unsafe_allow_html=True)

# --- INTERFAZ DE USUARIO ---
st.set_page_config(layout="wide")
st.title("📈 Historial de Actividades y Consejos")
start_date = st.date_input("Fecha de inicio", datetime.now().date() - timedelta(days=7))
end_date = st.date_input("Fecha de fin", datetime.now().date())

# --- LÓGICA PRINCIPAL ---
if start_date and end_date:
    if start_date > end_date:
        st.error("Error: La fecha de inicio no puede ser posterior a la fecha de fin.")
    else:
        params = {'oldest': start_date.strftime('%Y-%m-%d'), 'newest': end_date.strftime('%Y-%m-%d')}
        
        try:
            activities_url = f"https://intervals.icu/api/v1/athlete/{ATHLETE_ID}/activities"
            activities_response = requests.get(activities_url, auth=('API_KEY', API_KEY), params=params)
            activities_json = activities_response.json() if activities_response.status_code == 200 else []

            if activities_json:
                METRIC_MAP = {
                    'name': 'Actividad', 'start_date_local': 'Fecha', 'type': 'Tipo',
                    'moving_time': 'Duración', 'distance': 'Distancia (km)',
                    'icu_training_load': 'TSS', 'icu_intensity': 'Intensidad (%)',
                    'icu_weighted_avg_watts': 'Potencia Norm. (W)', 'average_watts': 'Potencia Media (W)',
                    'max_watts': 'Potencia Máx (W)', 'average_heartrate': 'FC Media',
                    'max_heartrate': 'FC Máx', 'average_cadence': 'Cadencia Media',
                    'icu_power_hr': 'Potencia/FC', 'icu_power_hr_z2': 'Potencia/FC en Z2', 
                    'icu_decoupling': 'Desacople (Pw:HR)', 'icu_ctl': 'CTL', 'icu_atl': 'ATL'
                }
                
                wellness_url = f"https://intervals.icu/api/v1/athlete/{ATHLETE_ID}/wellness"
                wellness_response = requests.get(wellness_url, auth=('API_KEY', API_KEY), params=params)
                wellness_json = wellness_response.json() if wellness_response.status_code == 200 else []
                vo2max_map = {item['id']: item.get('vo2max') for item in wellness_json if 'vo2max' in item and item.get('vo2max') is not None}

                all_available_metrics = list(METRIC_MAP.values()) + ['Eficiencia (NP/FC)', 'TSB', 'VO2max (Garmin)']
                processed_activities = []
                for activity in reversed(activities_json):
                    activity_data = {}
                    for api_key, value in activity.items():
                        if api_key in METRIC_MAP and value is not None and value != '':
                            human_name = METRIC_MAP[api_key]
                            if api_key == 'start_date_local': value = value[:10]
                            elif api_key == 'moving_time': value = format_duration(value)
                            elif api_key == 'distance': value = value / 1000
                            activity_data[human_name] = value
                    
                    np = activity.get('icu_weighted_avg_watts')
                    hr = activity.get('average_heartrate')
                    if np and isinstance(np, (int, float)) and hr and isinstance(hr, (int, float)) and hr > 0:
                        activity_data['Eficiencia (NP/FC)'] = np / hr
                    
                    activity_date_str = activity_data.get('Fecha')
                    if activity_date_str and activity_date_str in vo2max_map:
                        activity_data['VO2max (Garmin)'] = vo2max_map[activity_date_str]

                    if activity_data:
                        processed_activities.append(activity_data)
                
                if processed_activities:
                    df_activities = pd.DataFrame(processed_activities)
                    if 'Fecha' in df_activities.columns:
                        df_activities['Fecha'] = pd.to_datetime(df_activities['Fecha'])
                        df_activities = df_activities.sort_values(by='Fecha', ascending=False)
                    if 'CTL' in df_activities.columns and 'ATL' in df_activities.columns:
                        df_activities['TSB'] = df_activities['CTL'] - df_activities['ATL']

                    st.header("🚴 Resumen de Actividades")
                    
                    # Usamos df_activities.columns para asegurar que solo mostramos columnas que existen
                    default_cols = ['Fecha', 'Actividad', 'Duración', 'TSS', 'Potencia Norm. (W)', 'FC Media', 'CTL', 'ATL', 'TSB', 'Eficiencia (NP/FC)', 'VO2max (Garmin)']
                    final_default_cols = [col for col in default_cols if col in df_activities.columns]

                    selected_metrics = st.multiselect('Selecciona las métricas a mostrar:', options=sorted(list(df_activities.columns)), default=final_default_cols)
                    
                    if selected_metrics:
                        display_df = df_activities[selected_metrics].copy()
                        format_dict = {}
                        integer_cols = ['TSS', 'Potencia Norm. (W)', 'Potencia Media (W)', 'Potencia Máx (W)', 'FC Media', 'FC Máx', 'Cadencia Media']
                        
                        # --- INICIO DE LA CORRECCIÓN ---
                        # Se ha cambiado pd.api_types por pd.api.types
                        for col in display_df.columns:
                            if col in integer_cols: format_dict[col] = '{:.0f}'
                            elif pd.api.types.is_float_dtype(display_df[col]): format_dict[col] = '{:.1f}'
                            if col == 'Fecha': format_dict[col] = '{:%d-%m-%Y}'
                        # --- FIN DE LA CORRECCIÓN ---
                            
                        st.dataframe(display_df.style.format(format_dict, na_rep='-'), use_container_width=True)

                        word_file = dataframe_to_word(display_df)
                        st.download_button(
                            label="📥 Descargar como Word (.docx)",
                            data=word_file,
                            file_name=f"resumen_actividades_{start_date.strftime('%Y%m%d')}_{end_date.strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                else:
                    st.info("ℹ️ No se encontraron actividades con datos relevantes.")
            else:
                st.warning("No se encontraron actividades en el rango de fechas seleccionado.")
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Error de conexión de red al obtener actividades: {e}")

        st.markdown("---")
        st.header(f"💗 Estado de Salud para el Último Día ({end_date.strftime('%d-%m-%Y')})")
        readiness = get_readiness_analysis_v3(end_date, API_KEY, ATHLETE_ID)
        if readiness and "error" not in readiness:
            st.subheader(readiness['verdict'])
            st.caption("Puntuación de Salud:")
            display_gauge(readiness['readiness_score'])
            with st.expander("🔍 Analiza tu puntuación en detalle"):
                for line in readiness['score_breakdown']:
                    st.info(line)
                st.markdown(f"**PUNTUACIÓN TOTAL: {readiness['readiness_score']}**")
        else:
            st.error(readiness.get("error", "No se pudo generar el análisis de salud."))